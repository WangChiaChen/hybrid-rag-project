"""Phase 5：自動生成 Word 分析報告（精緻排版版）

排版重點：
- 全文套用中文字型（微軟正黑體），標題／表格都設 eastAsia，避免 Word 用 Calibri 顯示中文。
- 封面式標題 + 副標 + 產出資訊，中間一條細分隔線。
- 指標表：深藍表頭白字、隔列淡底、數值右對齊、變化欄依台灣慣例上紅下綠。
- 問答摘要：逐題排版（Q 粗體、A 分段），沒有問答就退回單純敘述。
"""
import base64
import io
import re
from datetime import date

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# 色票沿用網頁前端那套（深藍主色、台灣紅漲綠跌）
NAVY = RGBColor(0x1E, 0x3A, 0x5F)
GREY = RGBColor(0x76, 0x7C, 0x87)
DARK = RGBColor(0x1A, 0x1D, 0x21)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
UP_RED = RGBColor(0xC0, 0x39, 0x2B)   # 正／漲
DOWN_GREEN = RGBColor(0x2E, 0x7D, 0x46)  # 負／跌
NAVY_HEX = "1E3A5F"
ROW_HEX = "F2F5F9"
CJK_FONT = "Microsoft JhengHei"


def _style_run(run, size=11, bold=False, color=None, font=CJK_FONT):
    """統一設定字型——關鍵是同時設 eastAsia，否則中文會 fallback 成 Calibri。"""
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)


def _para(doc, text="", size=11, bold=False, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        _style_run(p.add_run(text), size=size, bold=bold, color=color)
    return p


def _rule(doc, color=NAVY_HEX, size="8"):
    """一條細分隔線（用段落的下框線做的）。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    p._p.get_or_add_pPr().append(pbdr)
    return p


def _section_head(doc, text):
    """章節標題：深藍粗體 + 下方細線。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    _style_run(p.add_run(text), size=14, bold=True, color=NAVY)
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "C9D6E4")
    pbdr.append(bottom)
    p._p.get_or_add_pPr().append(pbdr)


def _shade(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    cell._element.get_or_add_tcPr().append(shd)


def _cell(cell, text, size=10, bold=False, color=None, align=None, shade=None):
    para = cell.paragraphs[0]
    for r in list(para.runs):  # 清掉預設空 run，避免留下沒樣式的空白 run
        r._element.getparent().remove(r._element)
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.space_before = Pt(2)
    if align is not None:
        para.alignment = align
    _style_run(para.add_run(str(text)), size=size, bold=bold, color=color)
    if shade:
        _shade(cell, shade)


def _change_display(change):
    """把變化率轉成（文字, 顏色）。台灣慣例：正／漲＝紅，負／跌＝綠。"""
    if change is None or change == "":
        return "—", GREY
    try:
        v = float(str(change).replace("%", "").replace(",", ""))
    except ValueError:
        return str(change), DARK
    sign = "+" if v > 0 else ""
    color = UP_RED if v > 0 else DOWN_GREEN if v < 0 else GREY
    return f"{sign}{v}%", color


# ---------- 把 LLM 答案裡的 Markdown 轉成真正的 Word 排版 ----------
# LLM／EAP 的答案常帶 **粗體**、| 表格 |、- 條列。直接塞進 Word 會露出米字號跟管線符號，
# 表格也只是一串純文字。這裡做最小可用的 Markdown → docx 轉換。
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_TABLE_LINE = re.compile(r"^\s*\|.*\|?\s*$")
_BULLET_RE = re.compile(r"^\s*[-*]\s+")
_SEP_CELL = re.compile(r"^:?-{2,}:?$")


def _add_inline(paragraph, text, size=11, color=DARK, base_bold=False):
    """把一行文字加進段落，`**...**` 的部分變粗體，其餘照 base_bold。"""
    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            _style_run(paragraph.add_run(text[pos:m.start()]), size=size, bold=base_bold, color=color)
        _style_run(paragraph.add_run(m.group(1)), size=size, bold=True, color=color)
        pos = m.end()
    if pos < len(text):
        _style_run(paragraph.add_run(text[pos:]), size=size, bold=base_bold, color=color)


def _split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def _add_md_table(doc, block):
    rows = [_split_row(r) for r in block]
    is_sep = lambda cells: all(_SEP_CELL.match(c) or c == "" for c in cells)
    body = [r for r in rows if not is_sep(r)]
    if not body:
        return
    ncol = max(len(r) for r in body)
    table = doc.add_table(rows=0, cols=ncol)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for ridx, r in enumerate(body):
        cells = table.add_row().cells
        for c in range(ncol):
            txt = (r[c] if c < len(r) else "").replace("**", "")  # 表格內不再處理粗體，去掉米字號即可
            if ridx == 0:
                _cell(cells[c], txt, size=10, bold=True, color=WHITE,
                      align=WD_ALIGN_PARAGRAPH.CENTER, shade=NAVY_HEX)
            else:
                _cell(cells[c], txt, size=10, color=DARK,
                      align=WD_ALIGN_PARAGRAPH.LEFT, shade=ROW_HEX if ridx % 2 == 0 else None)


def _add_image(doc, data_url):
    """把前端傳來的 base64 PNG（趨勢圖）嵌進報告，置中。"""
    try:
        b64 = data_url.split(",", 1)[1] if "," in data_url else data_url
        raw = base64.b64decode(b64)
        doc.add_picture(io.BytesIO(raw), width=Inches(5.6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass  # 圖壞了不該讓整份報告產不出來


def _add_answer(doc, text):
    """把一段（可能含 Markdown 的）答案渲染成 Word：段落、粗體、表格、條列。"""
    lines = str(text).replace("\r", "").split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        if _TABLE_LINE.match(line) and "|" in line:      # 表格區塊
            block = []
            while i < len(lines) and _TABLE_LINE.match(lines[i]) and "|" in lines[i]:
                block.append(lines[i])
                i += 1
            _add_md_table(doc, block)
            continue
        if _BULLET_RE.match(line):                        # 條列
            while i < len(lines) and _BULLET_RE.match(lines[i]):
                p = _para(doc, "", size=11, space_after=3)
                p.paragraph_format.left_indent = Pt(14)
                _add_inline(p, "•　" + _BULLET_RE.sub("", lines[i]))
                i += 1
            continue
        hm = re.match(r"^\s*#{1,6}\s+(.*)$", line)         # 標題 #／##／### → 粗體小標
        if hm:
            p = _para(doc, "", size=12, space_after=3)
            p.paragraph_format.space_before = Pt(6)
            _style_run(p.add_run(hm.group(1).strip()), size=12, bold=True, color=NAVY)
            i += 1
            continue
        if not line.strip():                              # 空行
            i += 1
            continue
        p = _para(doc, "", size=11, space_after=6)        # 一般段落
        _add_inline(p, line)
        i += 1


def generate_report(company, period, metrics_summary, narrative_summary, output,
                    last_period=None, qa_pairs=None):
    """output 可以是檔案路徑，也可以是 BytesIO——python-docx 兩種都吃。
    傳檔案物件就不落地到伺服器磁碟，前端才能直接讓使用者下載。
    """
    doc = Document()
    # 讓內文預設字型就是中文字型（含 eastAsia），之後每個 run 也會再保險設一次
    normal = doc.styles["Normal"]
    normal.font.name = CJK_FONT
    normal.font.size = Pt(11)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)

    # ---------- 封面標題 ----------
    _para(doc, f"{company}　{period}", size=13, bold=True, color=GREY,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    _para(doc, "財務分析報告", size=24, bold=True, color=NAVY,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _para(doc, "Hybrid RAG · Vector RAG（語意檢索）＋ 結構化指標庫（公式計算）＋ AI Agent 路由",
          size=10, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

    meta = f"產出日期：{date.today().isoformat()}　｜　分析期間：{period}"
    if last_period:
        meta += f"　｜　比較基準：{last_period}"
    _para(doc, meta, size=9, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    _rule(doc)

    # ---------- 一、關鍵指標 ----------
    _section_head(doc, "一、關鍵指標變化")
    _para(doc, "", space_after=2)

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True

    headers = ["指標", "數值", "單位", f"變化（vs {last_period}）" if last_period else "變化"]
    aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT,
              WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT]
    for c, (h, al) in enumerate(zip(headers, aligns)):
        _cell(table.rows[0].cells[c], h, size=10, bold=True,
              color=RGBColor(0xFF, 0xFF, 0xFF), align=al, shade=NAVY_HEX)

    for i, m in enumerate(metrics_summary):
        cells = table.add_row().cells
        shade = ROW_HEX if i % 2 == 1 else None
        name = str(m.get("name", ""))
        if m.get("cumulative"):
            name += "（累計）"
        _cell(cells[0], name, size=10, color=DARK, align=WD_ALIGN_PARAGRAPH.LEFT, shade=shade)
        _cell(cells[1], m.get("value", ""), size=10, color=DARK, align=WD_ALIGN_PARAGRAPH.RIGHT, shade=shade)
        _cell(cells[2], m.get("unit", "") or "—", size=9, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, shade=shade)
        ctext, ccolor = _change_display(m.get("change"))
        _cell(cells[3], ctext, size=10, bold=(ctext != "—"), color=ccolor, align=WD_ALIGN_PARAGRAPH.RIGHT, shade=shade)

    _para(doc, "", space_after=2)
    # 單位帶「＊」代表原始簡報沒標、由系統依同公司其他期間或指標型別推定，要講清楚，
    # 讀者才知道哪些單位是揭露值、哪些是推定值。
    note = ("註：變化欄採台灣慣例，紅色為上升、綠色為下降。標「累計」者為年初至今累計值，僅同季跨年度可比。"
            "單位標「＊」者為系統推定（原始簡報未標示）；顯示「—」代表無從推定，請勿與其他數字直接比大小。")
    _para(doc, note, size=8, color=GREY, space_after=4)

    # ---------- 二、AI 問答摘要 ----------
    if qa_pairs:
        _section_head(doc, "二、AI 問答摘要")
        _para(doc, "以下為本次分析對話中，AI Agent 依真實財報數據給出的問答紀錄。",
              size=9, color=GREY, space_after=8)
        for idx, qa in enumerate(qa_pairs, 1):
            q = _para(doc, "", space_after=3)
            _style_run(q.add_run(f"Q{idx}　"), size=11, bold=True, color=NAVY)
            _add_inline(q, str(qa.get("question", "")), size=11, color=DARK, base_bold=True)
            a = _para(doc, "", space_after=2)
            _style_run(a.add_run("A"), size=11, bold=True, color=GREY)
            # 答案用 Markdown 渲染：粗體、表格、條列都變成真正的 Word 排版
            _add_answer(doc, str(qa.get("answer", "")).strip())
            if qa.get("image"):
                _add_image(doc, qa["image"])  # 該題的趨勢圖
            _para(doc, "", space_after=6)      # 題與題之間留白
    else:
        _section_head(doc, "二、經理人解釋摘要")
        _para(doc, narrative_summary or "（無相關敘述資料）", size=11, color=DARK, space_after=6)

    # ---------- 頁尾 ----------
    _rule(doc, color="DDDDDD", size="4")
    _para(doc,
          "資料來源：公開資訊觀測站 · 各公司投資人關係網站。數值由指標庫以公式計算，"
          "敘述由 AI Agent 依檢索到的法說會內容生成，僅供分析參考，不構成投資建議。",
          size=8, color=GREY, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=0)

    doc.save(output)


if __name__ == "__main__":
    generate_report(
        company="中信金控",
        period="2026Q1",
        last_period="2025Q4",
        metrics_summary=[
            {"name": "手續費淨收益", "value": "17,977", "unit": "百萬元", "change": "3.19", "cumulative": False},
            {"name": "基本每股盈餘", "value": "1.18", "unit": "元", "change": "-2.5", "cumulative": False},
            {"name": "股東權益報酬率(ROE)", "value": "17.46", "unit": "%", "change": "", "cumulative": False},
        ],
        narrative_summary="手續費淨收益成長主要受惠於財富管理業務回升。",
        qa_pairs=[
            {"question": "手續費淨收益變化多少？",
             "answer": "中信金控 2026Q1 的手續費淨收益為 **17,977 百萬元**，較 2025Q4 成長 3.19%。"},
            {"question": "跟玉山比呢？",
             "answer": "以下是兩家最近一季的比較：\n\n"
                       "| 指標 | 中信金控 | 玉山金控 | 哪家較佳 |\n"
                       "|------|---------|---------|---------|\n"
                       "| 稅後淨利 | 231.04 億元 | 100.57 億元 | 中信金控 |\n"
                       "| 每股盈餘（EPS） | 1.02 元 | 0.62 元 | 中信金控 |\n\n"
                       "重點說明：\n"
                       "- **獲利規模**：中信金控明顯較大。\n"
                       "- **成長動能**：玉山金控近期較快。"},
        ],
        output="outputs/sample_report.docx",
    )
    print("報告已生成：outputs/sample_report.docx")
